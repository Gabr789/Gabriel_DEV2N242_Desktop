# pip install python-docx
# Docx Viewer
import os
from docx import Document

class GerenciadorWord:

    def __init__(self, base_dir=None):
        self.base_dir = base_dir or os.getcwd()

    def definir_diretorio_base(self, caminho):
        if os.path.isdir(caminho):
            self.base_dir = caminho
        else:
            raise ValueError(f"Diretório inválido: {caminho}")

    def _caminho(self, nome):
        return os.path.join(self.base_dir, nome)

    def ler_arquivo(self, nome):
        caminho = self._caminho(nome)
        if not os.path.isfile(caminho):
            raise FileNotFoundError(f"Arquivo não encontrado: {caminho}")
        doc = Document(caminho)
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)

    def escrever_arquivo(self, nome, texto, modo='w'):
        caminho = self._caminho(nome)
        doc = Document()
        for linha in texto.split('\n'):
            doc.add_paragraph(linha)
        doc.save(caminho)

    def existe(self, nome):
        return os.path.isfile(self._caminho(nome))

    def remover_arquivo(self, nome):
        caminho = self._caminho(nome)
        if os.path.isfile(caminho):
            os.remove(caminho)

    def renomear(self, antigo, novo):
        os.rename(self._caminho(antigo), self._caminho(novo))

    def contar_linhas(self, nome):
        texto = self.ler_arquivo(nome)
        return len(texto.split('\n'))

    def contar_palavras(self, nome):
        texto = self.ler_arquivo(nome)
        return len(texto.split())

    def substituir_texto(self, nome, texto_antigo, texto_novo):
        caminho = self._caminho(nome)
        doc = Document(caminho)
        for paragraph in doc.paragraphs:
            if texto_antigo in paragraph.text:
                paragraph.text = paragraph.text.replace(texto_antigo, texto_novo)
        doc.save(caminho)

    def copiar_arquivo(self, origem, destino):
        caminho_origem = self._caminho(origem)
        caminho_destino = self._caminho(destino)
        with open(caminho_origem, 'rb') as fsrc:
            with open(caminho_destino, 'wb') as fdst:
                fdst.write(fsrc.read())

    def criar_arquivo_teste(self, nome="arquivo_teste.docx"):
        self.escrever_arquivo(nome, "Este é um arquivo de teste.")